import tempfile
import os
import sys
from flask import Flask, render_template, request, jsonify, send_file
import assemblyai as aai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from waitress import serve

app = Flask(__name__)

# Set your AssemblyAI API key
aai.settings.api_key = "b9d5344d2aa744318723eee6c06be532"

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    files = request.files.getlist('file')
    user_api_key = request.form.get('apiKey')
    output_format = request.form.get('outputFormat')
    include_speaker_labels = request.form.get('speakerLabels') == 'yes'

    if not user_api_key:
        return jsonify({'error': 'AssemblyAI API key is required'}), 400

    aai.settings.api_key = user_api_key
    output_dir = tempfile.gettempdir()
    results = []

    for file in files:
        input_source = None
        try:
            if file:
                # Save each file to a temporary directory
                temp_dir = tempfile.gettempdir()
                file_path = os.path.join(temp_dir, file.filename)
                file.save(file_path)
                input_source = file_path

                # Configure the transcription with additional options
                transcription_config = aai.TranscriptionConfig(
                    speaker_labels=include_speaker_labels,
                    punctuate=True,
                    format_text=True,
                    language_detection=True  # Auto-detect the language
                )

                transcriber = aai.Transcriber()

                # Attempt to transcribe each file
                transcript = transcriber.transcribe(input_source, config=transcription_config)

                # Debugging: Print the entire transcript object for analysis
                print(transcript)

                # Check if the transcription process was successful
                if transcript.status != 'completed':
                    return jsonify({'error': f"Transcription failed with status: {transcript.status}"}), 500

                # Additional check for empty transcription
                if not transcript.utterances and not transcript.words:
                    return jsonify({'error': 'Transcription returned no results.'}), 500

                # Debugging output path
                print(f"Saving output to: {output_dir}")

                # Handle captions output based on user selection
                if output_format == "srt":
                    if transcript.words:  # Ensure there are words detected to create captions
                        srt_output_filename = os.path.join(output_dir, os.path.splitext(os.path.basename(input_source))[0] + ".srt")
                        print(f"Writing SRT file to: {srt_output_filename}")
                        with open(srt_output_filename, "w") as srt_file:
                            srt_file.write(transcript.export_subtitles_srt())
                        output_filename = srt_output_filename
                    else:
                        return jsonify({'error': 'No words detected for captions generation.'}), 500
                elif output_format == "vtt":
                    if transcript.words:
                        vtt_output_filename = os.path.join(output_dir, os.path.splitext(os.path.basename(input_source))[0] + ".vtt")
                        print(f"Writing VTT file to: {vtt_output_filename}")
                        with open(vtt_output_filename, "w") as vtt_file:
                            vtt_file.write(transcript.export_subtitles_vtt())
                        output_filename = vtt_output_filename
                    else:
                        return jsonify({'error': 'No words detected for captions generation.'}), 500
                else:  # Handle transcription output
                    output_filename = os.path.join(output_dir, os.path.splitext(os.path.basename(input_source))[0] + ".docx")
                    print(f"Writing DOCX file to: {output_filename}")
                    doc = Document()

                    # Set the document style to "No Spacing"
                    no_space_style = doc.styles['No Spacing']
                    font = no_space_style.font
                    font.name = 'Verdana'
                    font.size = Pt(11)

                    for utterance in transcript.utterances:
                        # Add speaker label in bold
                        speaker_paragraph = doc.add_paragraph(style='No Spacing')
                        speaker_run = speaker_paragraph.add_run(f"Speaker {utterance.speaker}")
                        speaker_run.bold = True
                        speaker_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        # Add speaker text in the next line
                        text_paragraph = doc.add_paragraph(utterance.text, style='No Spacing')
                        text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        # Add a blank paragraph for spacing between speaker sections
                        doc.add_paragraph()

                    # Add "END" at the bottom with a bottom border
                    end_paragraph = doc.add_paragraph(style='No Spacing')
                    end_run = end_paragraph.add_run("END")
                    end_run.bold = True
                    end_run.font.name = 'Verdana'
                    end_run.font.size = Pt(11)
                    end_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    # Add a bottom border under the "END"
                    for paragraph in doc.paragraphs:
                        if paragraph.text == "END":
                            p = paragraph._element
                            pPr = p.get_or_add_pPr()
                            borders = OxmlElement('w:pBdr')
                            bottom = OxmlElement('w:bottom')
                            bottom.set(qn('w:val'), 'single')
                            bottom.set(qn('w:sz'), '6')
                            bottom.set(qn('w:space'), '1')
                            bottom.set(qn('w:color'), '000000')
                            borders.append(bottom)
                            pPr.append(borders)

                    # Save the document
                    doc.save(output_filename)

        except Exception as e:
            return jsonify({'error': f'Transcription failed due to an error: {str(e)}'}), 500

        finally:
            if input_source:
                os.remove(input_source)

    return send_file(output_filename, as_attachment=True)

if __name__ == "__main__":
    serve(app, host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
